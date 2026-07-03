#!/usr/bin/env python3
"""
EEW COMP3B — DOCX → HTML (full auto)
Legge tutti i dati dal docx e genera la dashboard aggiornata.

Uso:
  python auto_report.py IR-076-EEW_COMP3B_-_18_05_2026_-_24_05_2026.docx
  python auto_report.py report.docx --output ./reports
  python auto_report.py report.docx --push
"""

import os, sys, re, base64, subprocess, zipfile, json
from datetime import datetime, timedelta
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("❌ pip install python-docx"); sys.exit(1)

# ── Calibrazione IR number ─────────────────────────────────────────────────
BASE_WEEK, BASE_IR = 19, 74

# ══════════════════════════════════════════════════════════════════════════════
# PARSER — legge tutto dal docx
# ══════════════════════════════════════════════════════════════════════════════

class DocxParser:

    def __init__(self, path):
        self.path = path
        self.doc  = Document(path)
        self.tables = self.doc.tables
        self.paras  = [p.text.strip() for p in self.doc.paragraphs]
        self.data   = {}
        self._parse()

    # ── helpers ───────────────────────────────────────────────────────────────

    def _tbl(self, idx):
        """Restituisce tabella come lista di righe (lista di celle).

        Le celle unite (merge) vengono collassate per IDENTITÀ dell'elemento XML
        (stesso <w:tc>), non per valore di testo: così due celle vuote distinte
        restano due colonne separate e l'allineamento delle colonne è corretto.
        """
        t = self.tables[idx]
        rows = []
        for row in t.rows:
            out = []
            seen = set()
            for c in row.cells:
                tc = id(c._tc)
                if tc in seen:
                    continue
                seen.add(tc)
                out.append(c.text.strip())
            rows.append(out)
        return rows

    def _find_tbl(self, *keywords):
        """Trova la prima tabella che contiene tutte le keyword nell'header."""
        for i, t in enumerate(self.tables):
            if not t.rows: continue
            header = ' '.join(c.text for c in t.rows[0].cells).upper()
            if all(k.upper() in header for k in keywords):
                return i
        return -1

    def _para_after(self, keyword, n=5):
        """Restituisce i paragrafi dopo una keyword."""
        idx = next((i for i,p in enumerate(self.paras) if keyword.upper() in p.upper()), -1)
        if idx < 0: return []
        return self.paras[idx:idx+n]

    def _num(self, s):
        """Estrae numero da stringa (gestisce , come sep decimale)."""
        s = str(s).replace('.','').replace(',','.')
        m = re.search(r'\d[\d.]*', s)
        return float(m.group()) if m else 0.0

    def _pct(self, s):
        """Estrae percentuale."""
        m = re.search(r'(\d+)[,.](\d+)%?', str(s))
        if m: return f"{m.group(1)},{m.group(2)}%"
        m2 = re.search(r'(\d+)%', str(s))
        return m2.group(0) if m2 else s

    # ── parse header (tabella 0) ───────────────────────────────────────────────

    def _parse_header(self):
        rows = self._tbl(0)
        d = self.data
        for row in rows:
            text = ' '.join(row)
            # Date visita
            m = re.search(r'(\d{2}\.\d{2}\.\d{4})', text)
            if 'From:' in text or 'Date of visit' in text:
                dates = re.findall(r'(\d{2})\.(\d{2})\.(\d{4})', text)
                if len(dates) >= 2:
                    d1,m1,y1 = dates[0]; d2,m2,y2 = dates[1]
                    d['week_start'] = datetime(int(y1),int(m1),int(d1))
                    d['week_end']   = datetime(int(y2),int(m2),int(d2))
            # Reference Week
            if 'Reference Week:' in text or 'Reference Week' in ' '.join(row):
                for cell in row:
                    if re.match(r'^\d{1,2}$', cell.strip()):
                        d['week_number'] = int(cell.strip())

        # Fallback da nome file
        if 'week_start' not in d:
            stem = Path(self.path).stem
            m = re.search(r'(\d{2})_(\d{2})_(\d{4})_-_(\d{2})_(\d{2})_(\d{4})', stem)
            if m:
                dd1,mm1,yy1,dd2,mm2,yy2 = m.groups()
                d['week_start'] = datetime(int(yy1),int(mm1),int(dd1))
                d['week_end']   = datetime(int(yy2),int(mm2),int(dd2))
            else:
                today = datetime.now()
                d['week_start'] = today - timedelta(days=today.weekday())
                d['week_end']   = d['week_start'] + timedelta(days=6)

        if 'week_number' not in d:
            d['week_number'] = d['week_start'].isocalendar()[1]

        d['ir_number'] = BASE_IR + (d['week_number'] - BASE_WEEK)

    # ── parse contatti (tabella 2) ────────────────────────────────────────────

    def _parse_contacts(self):
        ti = self._find_tbl('Name','Company','Tasks')
        if ti < 0: return
        contacts = []
        for row in self._tbl(ti)[1:]:
            if len(row) >= 2 and row[0] and row[1]:
                contacts.append({
                    'name':    row[0] if len(row)>0 else '',
                    'company': row[1] if len(row)>1 else '',
                    'role':    row[2] if len(row)>2 else '',
                    'email':   row[3] if len(row)>3 else '',
                })
        self.data['contacts'] = contacts

    # ── parse dispatch (tabelle 9, 10, 11, 12) ───────────────────────────────

    def _parse_dispatch(self):
        d = self.data

        # PO summary (tabella 10) — niente riga "Total": i totali si sommano
        ti = self._find_tbl('SAIPEM PO','Pipes Released')
        if ti >= 0:
            rows = self._tbl(ti)
            po_data = []
            tot_rel = tot_disp = tot_ready = 0
            for row in rows[1:]:
                if row[0] in ('Total.','Total'):
                    continue
                if row[0].startswith('15'):
                    po_data.append({
                        'po':         row[0],
                        'released':   row[1] if len(row)>1 else '—',
                        'length_mm':  row[2] if len(row)>2 else '—',
                        'dispatched': row[3] if len(row)>3 else '—',
                        'ready':      row[4] if len(row)>4 else '—',
                    })
                    tot_rel   += int(self._num(row[1])) if len(row)>1 else 0
                    tot_disp  += int(self._num(row[3])) if len(row)>3 else 0
                    tot_ready += int(self._num(row[4])) if len(row)>4 else 0
            d['dispatch_po'] = po_data
            d['dispatch_total'] = {
                'released':   f"{tot_rel:,}".replace(',', '.'),
                'length_mm':  '—',
                'dispatched': f"{tot_disp:,}".replace(',', '.'),
                'ready':      f"{tot_ready:,}".replace(',', '.'),
            }

        # Dettaglio per categoria (tabella 10)
        ti2 = self._find_tbl('SAIPEM Pos','NOS Released','ACTUAL LENGTH')
        if ti2 >= 0:
            rows = self._tbl(ti2)
            cats = []
            for row in rows[1:]:
                if row[0] and not row[0].startswith('15') and row[0] != 'Total.':
                    cats.append({'cat': row[0], 'released': row[1] if len(row)>1 else '—',
                                 'length': row[2] if len(row)>2 else '—',
                                 'dispatched': row[3] if len(row)>3 else '—',
                                 'ready': row[4] if len(row)>4 else '—'})
            d['dispatch_cats'] = cats

        # Antwerp (tabella 11)
        ti3 = self._find_tbl('Stacked at Antwerp')
        if ti3 >= 0:
            rows = self._tbl(ti3)
            antwerp = []
            for row in rows[1:]:
                antwerp.append({'sow': row[0], 'nos': row[1], 'length': row[2]})
            d['antwerp'] = antwerp

        # Inspected (tabella 9)
        ti4 = self._find_tbl('SAIPEM PO','INSPECTED','LENGTH')
        if ti4 >= 0:
            rows = self._tbl(ti4)
            insp = []
            for row in rows[1:]:
                if row[0] in ('Total.','Total'):
                    d['inspected_total'] = row[1] if len(row)>1 else '—'
                    d['inspected_length'] = row[2] if len(row)>2 else '—'
                elif row[0].startswith('15'):
                    insp.append({'po': row[0],
                                 'nos': row[1] if len(row)>1 else '—',
                                 'length': row[2] if len(row)>2 else '—'})
            d['inspected_po'] = insp

    # ── parse fasi produzione (tabella 13) ───────────────────────────────────

    def _parse_phases(self):
        ti = self._find_tbl('Phase','Completed','Remaining','Total')
        if ti < 0: return
        rows = self._tbl(ti)
        phases = []
        for row in rows[1:]:
            if len(row) >= 4 and row[0]:
                phases.append({
                    'name':      row[0],
                    'completed': row[1],
                    'remaining': row[2],
                    'total':     row[3],
                    'pct':       row[4] if len(row)>4 else '',
                })
        self.data['phases'] = phases

    # ── parse materiali (tabella 14) ─────────────────────────────────────────

    def _parse_materials(self):
        ti = self._find_tbl('Material','Quantity','Percentage')
        if ti < 0: return
        rows = self._tbl(ti)
        self.data['materials'] = [
            {'type': r[0], 'qty': r[1], 'pct': r[2]}
            for r in rows[1:] if r[0]
        ]

    # ── parse item summary (tabella 16) ──────────────────────────────────────

    def _parse_items(self):
        ti = self._find_tbl('Item Nr.','Pipes Scheduled','Avg Completion','Status')
        if ti < 0: return
        rows = self._tbl(ti)
        self.data['items'] = [
            {'item': r[0], 'pipes': r[1], 'completion': r[2], 'status': r[3]}
            for r in rows[1:] if r[0]
        ]

    # ── parse yearly production (tabella 17) ─────────────────────────────────

    def _parse_yearly(self):
        ti = self._find_tbl('Year','Milling','Welding Base')
        if ti < 0: return
        self.data['yearly'] = self._tbl(ti)

    # ── parse WoW (tabella 18) ───────────────────────────────────────────────

    def _parse_wow(self):
        ti = self._find_tbl('WEEK-ON-WEEK')
        if ti < 0: return
        rows = self._tbl(ti)
        header = rows[0][0] if rows else ''
        m = re.search(r'W(\d+)\s*vs\s*W(\d+)', header, re.IGNORECASE)
        self.data['wow_label'] = header
        self.data['wow_prev']  = m.group(1) if m else ''
        self.data['wow_curr']  = m.group(2) if m else ''
        self.data['wow_rows']  = rows[1:]

    # ── parse attività settimanali (tabella 19) ──────────────────────────────

    def _parse_activities(self):
        # Colonne attuali: Date | FR No. | ITP Pos. | Activity Description | Pipe Nos | Result
        ti = self._find_tbl('Date','ITP Pos','Activity Description')
        if ti < 0:
            ti = self._find_tbl('Date','ITP Pos.','Description')
        if ti < 0:
            ti = self._find_tbl('Date','ITP','Description')
        if ti < 0: return
        rows = self._tbl(ti)
        last_date = ''
        acts = []
        for r in rows[1:]:
            if len(r) < 4 or not r[2]:   # serve almeno l'ITP Pos.
                continue
            date = r[0] if len(r)>0 and r[0] else last_date
            last_date = date
            acts.append({
                'date':   date,
                'fr':     r[1] if len(r)>1 else '',
                'itp':    r[2] if len(r)>2 else '',
                'desc':   r[3] if len(r)>3 else '',
                'pipes':  r[4] if len(r)>4 else '',
                'result': r[5] if len(r)>5 else '',
            })
        self.data['activities'] = acts

    # ── parse VAGB plates (tabelle 20, 21, 22) ───────────────────────────────

    def _parse_plates(self):
        # KPI (tabella 21)
        ti = self._find_tbl('KPI','Current Value','Target','Delta')
        if ti >= 0:
            rows = self._tbl(ti)
            kpis = {r[0]: {'value': r[1] if len(r)>1 else '',
                           'target': r[2] if len(r)>2 else '',
                           'delta': r[3] if len(r)>3 else ''}
                    for r in rows[1:] if r[0]}
            self.data['vagb_kpi'] = kpis
            # Estrai valori chiave
            for k, v in kpis.items():
                if 'Plates at EEW' in k and 'VAGB' in k:
                    self.data['vagb_pct']  = v['value']
                    self.data['vagb_delta'] = v['delta']
                if 'Number Plates at EEW' in k:
                    self.data['vagb_nos'] = v['value']
                if 'Total Plates Ordered' in k:
                    self.data['vagb_ordered'] = v['value']

        # Delivery summary (tabella 20)
        ti2 = self._find_tbl('Item','Total Plates','Delivered','Pending','Ordered')
        if ti2 >= 0:
            rows = self._tbl(ti2)
            self.data['vagb_delivery'] = [
                {'item': r[0] if len(r)>0 else '', 'total': r[1] if len(r)>1 else '',
                 'delivered': r[2] if len(r)>2 else '', 'pending': r[3] if len(r)>3 else '',
                 'ordered': r[4] if len(r)>4 else '', 'pct': r[5] if len(r)>5 else '',
                 'first': r[6] if len(r)>6 else '', 'last': r[7] if len(r)>7 else ''}
                for r in rows[1:] if r[0] and not r[0].startswith('TOTAL')
            ]

        # Schedule (tabella 22)
        ti3 = self._find_tbl('Calendar Week','Dispatch','Ex-Works')
        if ti3 >= 0:
            rows = self._tbl(ti3)
            self.data['vagb_schedule'] = [
                {'week': r[0] if len(r)>0 else '', 'dispatch': r[1] if len(r)>1 else '',
                 'schedule': r[2] if len(r)>2 else '',
                 'received': r[3] if len(r)>3 else ''}
                for r in rows[1:] if r[0]
            ]

    # ── parse HSE (tabella 32) ───────────────────────────────────────────────

    def _parse_hse(self):
        ti = self._find_tbl('No.','Description','Yes','No')
        if ti < 0: return
        rows = self._tbl(ti)
        hse = {}
        for row in rows[1:]:
            if row[0] in ('1','2','3','4'):
                hse[row[0]] = {'desc': row[1] if len(row)>1 else '',
                               'yes': bool(len(row)>2 and row[2] and row[2].lower() not in ('','-')),
                               'no':  bool(len(row)>3 and row[3] and row[3].lower() not in ('','-'))}
        self.data['hse'] = hse
        # Nota finale
        for row in rows:
            if 'No HSE issues' in ' '.join(row) or 'Nothing to report' in ' '.join(row):
                self.data['hse_note'] = row[1] if len(row)>1 else row[0]
                break

    # ── parse foto (tabella 36 + immagini) ───────────────────────────────────

    def _parse_photos(self):
        """Legge le foto e le caption dalla tabella foto e dalle relazioni del docx."""
        # Caption dalla tabella
        ti = -1
        for i, t in enumerate(self.tables):
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            flat = ' '.join(c for r in rows for c in r)
            if 'ITP STEP' in flat.upper() and 'PHOTO' in ' '.join(
                    [self.tables[i-1].rows[0].cells[0].text if i>0 else '']).upper():
                ti = i; break
        # Fallback: cerca tabella con "ITP STEP" nei contenuti
        if ti < 0:
            for i, t in enumerate(self.tables):
                flat = ' '.join(c.text for r in t.rows for c in r.cells)
                if re.search(r'ITP STEP \d+', flat):
                    ti = i; break

        # Captions in ordine di tabella
        captions_ordered = []
        if ti >= 0:
            for row in self.tables[ti].rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and 'ITP STEP' in txt.upper() and txt not in captions_ordered:
                        captions_ordered.append(txt)

        # Estrai immagini JPEG dal zip
        with zipfile.ZipFile(self.path) as z:
            rels_xml = z.read('word/_rels/document.xml.rels').decode()
            doc_xml  = z.read('word/document.xml').decode('utf-8', errors='ignore')

        # Map rId → filename
        rid_map = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/(image\d+\.(jpeg|jpg|png))"',
                             rels_xml, re.IGNORECASE):
            rid_map[m.group(1)] = m.group(2)

        # Trova rId delle immagini nell'ordine del documento
        rids_in_order = re.findall(r'r:embed="(rId\d+)"', doc_xml)
        seen = set()
        photo_rids = []
        for rid in rids_in_order:
            if rid in rid_map and rid not in seen:
                seen.add(rid)
                photo_rids.append(rid)

        # Leggi blob
        photos = []
        with zipfile.ZipFile(self.path) as z:
            for i, rid in enumerate(photo_rids):
                fname = rid_map[rid]
                # Salta EMF e immagini troppo piccole (loghi <20KB)
                if not fname.lower().endswith(('.jpeg','.jpg','.png')):
                    continue
                try:
                    blob = z.read(f'word/media/{fname}')
                    if len(blob) < 15000:   # salta loghi piccoli
                        continue
                    b64     = base64.b64encode(blob).decode()
                    caption = captions_ordered[len(photos)] if len(photos) < len(captions_ordered) \
                              else f'Photo {len(photos)+1}'
                    # Normalizza caption
                    caption = caption.title().replace('Itp','ITP').replace('Dpi','DPI')
                    ext = 'jpeg' if fname.lower().endswith(('.jpg','.jpeg')) else 'png'
                    photos.append({'caption': caption, 'b64': b64, 'ext': ext})
                except Exception:
                    continue

        self.data['photos'] = photos
        print(f"  📷 {len(photos)} foto estratte")

    # ── parse AOB (paragrafi) ─────────────────────────────────────────────────

    def _parse_aob(self):
        paras = self.paras
        aob_start = next((i for i,p in enumerate(paras)
                          if 'AREA OF CONCERN' in p.upper()), -1)
        if aob_start < 0:
            self.data['aob'] = []; return
        # Raccogli paragrafi AOB fino a SECTION o fine
        aob = []
        for p in paras[aob_start:]:
            if re.match(r'SECTION (III|IV|V|VI)', p.upper()): break
            if p and not p.startswith('INTERACTIVE'):
                aob.append(p)
        self.data['aob'] = aob

    # ── parse enclosures (tabella 38) ────────────────────────────────────────

    def _parse_enclosures(self):
        ti = self._find_tbl('Document No.','Title - Description','Doc. Type')
        if ti < 0: return
        rows = self._tbl(ti)
        self.data['enclosures'] = [
            {'num': r[0], 'title': r[1], 'type': r[2]}
            for r in rows[1:] if r[0] and r[0].isdigit()
        ]

    # ── parse tally lists (tabella dopo activities) ───────────────────────────

    def _parse_tally(self):
        ti = self._find_tbl('Tally List No.','Quantity','Length')
        if ti < 0:
            ti = self._find_tbl('Tally List','Item','Quantity')
        if ti < 0: return
        rows = self._tbl(ti)
        tally = []
        for r in rows[1:]:
            if not r[0]: continue
            if r[0].upper() in ('TOTAL', 'TOTAL.'): continue
            tally.append({
                'date':     r[0] if len(r)>0 else '',
                'no':       r[1] if len(r)>1 else '',
                'item':     r[2] if len(r)>2 else '',
                'qty':      r[3] if len(r)>3 else '',
                'length':   r[4] if len(r)>4 else '',
                'saipem':   r[5] if len(r)>5 else '',
                'remarks':  r[6] if len(r)>6 else '',
            })
        # leggi anche riga TOTAL
        for r in rows[1:]:
            if r[0].upper() in ('TOTAL', 'TOTAL.'):
                self.data['tally_total'] = {'qty': r[3] if len(r)>3 else '', 'length': r[4] if len(r)>4 else ''}
                break
        self.data['tally'] = tally

    # ── parse completion forecast by item & phase (tabella 15) ────────────────

    def _parse_forecast(self):
        ti = self._find_tbl('Phase','Actual','Target','Est. Completion')
        if ti < 0: return
        rows = self._tbl(ti)
        last_item = ''
        out = []
        for r in rows[1:]:
            if len(r) < 5 or not r[1]:   # serve la fase
                continue
            item = r[0] if r[0] else last_item
            last_item = item
            out.append({
                'item':   item,
                'phase':  r[1],
                'actual': r[2] if len(r)>2 else '',
                'target': r[3] if len(r)>3 else '',
                'remaining': r[4] if len(r)>4 else '',
                'rate':   r[5] if len(r)>5 else '',
                'prod7':  r[6] if len(r)>6 else '',
                'rate_req_d': r[7] if len(r)>7 else '',
                'est':    r[9] if len(r)>9 else '',
                'po_date': r[10] if len(r)>10 else '',
                'status': r[12] if len(r)>12 else (r[-1] if r else ''),
            })
        self.data['forecast'] = out

    # ── parse confronto con data PO (tabella 17) ──────────────────────────────

    def _parse_po_comparison(self):
        ti = self._find_tbl('COMPARISON WITH PO COMPLETION DATE')
        if ti < 0: return
        rows = self._tbl(ti)
        pairs = []
        for r in rows[1:]:
            if not r or 'COLOR LEGEND' in ' '.join(r).upper():
                break
            if len(r) >= 2 and r[0] and r[1]:
                pairs.append((r[0], r[1]))
        self.data['po_comparison'] = pairs

    # ── parse produzione per periodo (tabelle 11 settimana, 12 mese) ──────────

    def _parse_production(self):
        tw = self._find_tbl('Week','Milling','Welding Base')
        if tw >= 0:
            self.data['prod_week'] = self._tbl(tw)
        tm = self._find_tbl('Month','Milling','Welding Base')
        if tm >= 0:
            self.data['prod_month'] = self._tbl(tm)

    # ── parse prove di laboratorio / meccaniche (tabelle 20-25) ───────────────

    def _parse_mech(self):
        # Tabella riassuntiva principale (con esito)
        ti = self._find_tbl('Location','Pipe No.','Test Type','Result')
        if ti >= 0:
            self.data['mech'] = self._tbl(ti)
        # Dettagli opzionali
        tc = self._find_tbl('Test Type','Start Wt','End Wt','Temp')
        if tc >= 0:
            self.data['mech_corrosion'] = self._tbl(tc)
        td = self._find_tbl('Specimen No.','Temperature','Shear Area')
        if td >= 0:
            self.data['mech_dwtt'] = self._tbl(td)
        tt = self._find_tbl('Pipe No.','Sample No.','Min Req')
        if tt >= 0:
            self.data['mech_tensile'] = self._tbl(tt)

    # ── parse MDR / certificati piastre (tabella 29) ──────────────────────────

    def _parse_mdr(self):
        ti = self._find_tbl('COMPLETION PERCENTAGE','MISSING CERTIFICATES')
        if ti < 0: return
        rows = self._tbl(ti)
        d = self.data
        if len(rows) > 1 and len(rows[1]) >= 2:
            d['mdr_pct']     = rows[1][0]
            d['mdr_missing'] = rows[1][1]
        # Dettaglio stato certificati (riga header interna "Plate / Certificate Status")
        detail = []
        start = -1
        for i, r in enumerate(rows):
            if r and 'Certificate Status' in r[0]:
                start = i; break
        if start >= 0:
            for r in rows[start+1:]:
                if len(r) >= 3 and r[0]:
                    detail.append({'status': r[0], 'qty': r[1] if len(r)>1 else '',
                                   'pct': r[2] if len(r)>2 else '',
                                   'op': r[3] if len(r)>3 else ''})
        d['mdr_detail'] = detail

    # ── parse NOI prossima visita (tabella 27) ────────────────────────────────

    def _parse_noi(self):
        ti = self._find_tbl('NOI','Activity','Date','Location')
        if ti < 0: return
        rows = self._tbl(ti)
        out = []
        for r in rows[1:]:
            if len(r) >= 4 and r[1]:
                out.append({'noi': r[0], 'activity': r[1],
                            'date': r[3] if len(r)>3 else '',
                            'time': r[4] if len(r)>4 else '',
                            'loc':  r[5] if len(r)>5 else '',
                            'status': r[6] if len(r)>6 else ''})
        self.data['noi'] = out

    # ── parse documenti ITP usati (tabella 4) ─────────────────────────────────

    def _parse_docs(self):
        ti = self._find_tbl('Document No.','Title - Description','Rev.','Status')
        if ti < 0: return
        rows = self._tbl(ti)
        out = []
        for r in rows[1:]:
            if not r or not r[0]:
                continue
            if len(r) == 1:   # riga-titolo di categoria
                out.append({'section': r[0]})
            elif len(r) >= 4:
                out.append({'num': r[0], 'title': r[1], 'rev': r[2], 'status': r[3]})
        self.data['docs'] = out

    # ── entry point ───────────────────────────────────────────────────────────

    def _parse(self):
        print("  📄 Parsing docx...")
        self._parse_header()
        self._parse_contacts()
        self._parse_dispatch()
        self._parse_phases()
        self._parse_materials()
        self._parse_items()
        self._parse_yearly()
        self._parse_wow()
        self._parse_activities()
        self._parse_tally()
        self._parse_plates()
        self._parse_forecast()
        self._parse_po_comparison()
        self._parse_production()
        self._parse_mech()
        self._parse_mdr()
        self._parse_noi()
        self._parse_docs()
        self._parse_hse()
        self._parse_photos()
        self._parse_aob()
        self._parse_enclosures()
        d = self.data
        print(f"  ✓ W{d['week_number']}  {d['week_start'].strftime('%d.%m.%Y')}–{d['week_end'].strftime('%d.%m.%Y')}  IR-{d['ir_number']:03d}")
        print(f"  ✓ {len(d.get('phases',[]))} fasi  |  {len(d.get('activities',[]))} attività  |  {len(d.get('photos',[]))} foto")


# ══════════════════════════════════════════════════════════════════════════════
# LOGIN GATE — pagina di accesso con password (client-side, hash SHA-256)
# ══════════════════════════════════════════════════════════════════════════════
# Le password NON sono salvate in chiaro: solo il loro hash SHA-256 è nel file.
# Password utente: Comp3B-TgO6V5
# Password admin : Comp3B-Admin-JIyRPh
# (per cambiarle: calcola il nuovo hash con hashlib.sha256(pwd.encode()).hexdigest()
#  e sostituisci le due stringhe qui sotto)

LOGIN_GATE_CSS = """
/* LOGIN GATE */
#login-screen{position:fixed;inset:0;z-index:9999;background:var(--bg3);display:flex;align-items:center;justify-content:center;padding:20px}
.login-box{background:var(--bg);border:0.5px solid var(--border);border-radius:var(--rl);padding:40px 36px 32px;width:360px;max-width:100%;position:relative;overflow:hidden;box-shadow:0 10px 40px rgba(0,0,0,.14)}
.login-box::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;background:linear-gradient(90deg,var(--accent2),var(--accent))}
.login-title{font-size:15px;font-weight:500;color:var(--text);text-align:center;margin-bottom:3px}
.login-sub{font-size:10px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.07em;text-align:center;margin-bottom:26px}
.login-label{font-size:10px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.05em;margin-bottom:6px;display:block}
.login-input{width:100%;padding:9px 12px;background:var(--bg3);border:0.5px solid var(--border2);border-radius:var(--r);color:var(--text);font-size:13px;font-family:var(--font);outline:none;transition:border-color .15s;box-sizing:border-box;letter-spacing:1px}
.login-input:focus{border-color:var(--accent2)}
.login-btn{width:100%;margin-top:16px;padding:10px;background:linear-gradient(90deg,var(--accent2),var(--accent));border:none;border-radius:var(--r);font-family:var(--font);font-size:13px;font-weight:500;color:#fff;cursor:pointer;transition:opacity .15s}
.login-btn:hover{opacity:.88}
.login-error{color:var(--danger);font-size:11px;text-align:center;margin-top:10px;min-height:14px}
.login-footer{font-size:10px;color:var(--text3);text-align:center;margin-top:20px;opacity:.7}
.admin-visible{display:none}
body.admin-mode .admin-visible{display:inline-flex}
.admin-badge{background:var(--pu-bg);color:var(--pu-t);border:none;font-family:var(--font)}
"""

LOGIN_GATE_HTML = """
<div id="login-screen">
  <div class="login-box">
    <div class="login-title">NFPS COMP3B &mdash; Weekly Report</div>
    <div class="login-sub">IQS S.r.l. &middot; Restricted Access</div>
    <label class="login-label">Password</label>
    <input class="login-input" id="login-pwd" type="password" placeholder="&bull;&bull;&bull;&bull;&bull;&bull;&bull;&bull;&bull;&bull;&bull;&bull;" autocomplete="current-password">
    <button class="login-btn" onclick="doLogin()">Access Report</button>
    <div class="login-error" id="login-err"></div>
    <div class="login-footer">IQS Confidential &middot; Order 45650</div>
  </div>
</div>
"""

LOGIN_GATE_JS = """
<script>
const PWD_HASH       = ["2bc8ca85c10c52d3ee9ffc26f9aa1943","6ec9575d96b507d2380f33566d502b6d"].join("");
const PWD_HASH_ADMIN = ["0fa5c082b9b61ac3d99b7ef45664422a","fd5ea8116808d07768f1284a7d06dadc"].join("");
const SESSION_KEY  = "comp3b_auth";
const SESSION_ROLE = "comp3b_role";

async function sha256(m){
  const b = await crypto.subtle.digest("SHA-256", new TextEncoder().encode(m));
  return Array.from(new Uint8Array(b)).map(x => x.toString(16).padStart(2,"0")).join("");
}

async function doLogin(){
  const p = document.getElementById("login-pwd").value;
  const h = await sha256(p);
  if (h === PWD_HASH_ADMIN) {
    sessionStorage.setItem(SESSION_KEY, "1");
    sessionStorage.setItem(SESSION_ROLE, "admin");
    document.getElementById("login-screen").style.display = "none";
    applyRole("admin");
  } else if (h === PWD_HASH) {
    sessionStorage.setItem(SESSION_KEY, "1");
    sessionStorage.setItem(SESSION_ROLE, "user");
    document.getElementById("login-screen").style.display = "none";
    applyRole("user");
  } else {
    const e = document.getElementById("login-err");
    e.textContent = "Password errata \\u2014 riprova";
    document.getElementById("login-pwd").value = "";
    setTimeout(() => e.textContent = "", 3000);
  }
}

function applyRole(role){
  if (role === "admin") document.body.classList.add("admin-mode");
  else document.body.classList.remove("admin-mode");
}

document.addEventListener("DOMContentLoaded", () => {
  const i = document.getElementById("login-pwd");
  if (i) i.addEventListener("keydown", e => { if (e.key === "Enter") doLogin(); });
  const stored = sessionStorage.getItem(SESSION_KEY);
  const role   = sessionStorage.getItem(SESSION_ROLE) || "user";
  if (stored === "1") {
    document.getElementById("login-screen").style.display = "none";
    applyRole(role);
  }
});
</script>
"""

# ══════════════════════════════════════════════════════════════════════════════
# HTML BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def esc(s):
    """Escape HTML."""
    return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

def pill(text, cls='pg'):
    return f'<span class="pill {cls}">{esc(text)}</span>'

def status_pill(s):
    s = str(s).strip().upper()
    if s in ('COMPLETE','COMPLETED','AFC','CLOSED','ISSUED','YES'):
        return pill(s.title(), 'pg')
    if s in ('ONGOING','HALFWAY','IFR','IN PROGRESS'):
        return pill(s.title(), 'pa')
    if s in ('IFA',):
        return pill(s, 'pb')
    if s in ('NOT SENT','OPEN','NOT YET','NO'):
        return pill(s.title(), 'pr')
    return pill(s, 'pgr')

def build_html(d):
    ws = d['week_start']; we = d['week_end']
    wn = d['week_number']; ir = d['ir_number']

    dot_end    = we.strftime("%d.%m.%Y")
    slash_end  = we.strftime("%d/%m/%Y")
    slash_next = (we+timedelta(days=1)).strftime("%d/%m/%Y")
    range_en   = f"{ws.strftime('%d')} to {we.strftime('%d')} {we.strftime('%B %Y')}"
    sidebar_w  = f"W{wn} &bull; {ws.strftime('%d')}&#8211;{we.strftime('%d')} {we.strftime('%B %Y')}"

    # ── KPIs da dati parsed ───────────────────────────────────────────────────
    dt = d.get('dispatch_total', {})
    pipes_released  = dt.get('released',  '—')
    pipes_dispatched= dt.get('dispatched','—')
    pipes_ready     = dt.get('ready',     '—')
    pipes_inspected = d.get('inspected_total','—')
    vagb_pct        = d.get('vagb_pct',  '—')
    vagb_nos        = d.get('vagb_nos',  '—')
    vagb_ordered    = d.get('vagb_ordered','4,502')
    vagb_delta      = d.get('vagb_delta', '—')

    # Antwerp totale
    antwerp_total = '—'
    for row in d.get('antwerp',[]):
        if row['sow'].upper() in ('TOTAL','TOTAL.'):
            antwerp_total = row['nos']; break

    # Item 00400 progress
    item400_pct = '—'
    for it in d.get('items',[]):
        if it['item'] == '00400':
            item400_pct = f"{it['completion']}%"; break

    # ── Photos JS ─────────────────────────────────────────────────────────────
    photos = d.get('photos', [])
    photos_data = [[p['caption'], f"data:image/{p['ext']};base64,{p['b64']}"] for p in photos]
    photos_js = f"const PHOTOS = {json.dumps(photos_data)};"

    # Raggruppa foto per sezione (ITP STEP N → gruppo)
    photo_groups = {}
    for i, p in enumerate(photos):
        m = re.search(r'ITP Step (\d+)', p['caption'], re.IGNORECASE)
        step = m.group(1) if m else 'Other'
        photo_groups.setdefault(step, []).append(i)

    # ── Sezioni HTML ──────────────────────────────────────────────────────────

    # Overview alerts dal AOB
    aob = d.get('aob', [])
    alerts_html = ''
    for p in aob:
        if 'recovery' in p.lower() or 'Recovery' in p:
            alerts_html += f'<div class="alert ag"><strong>VAGB recovery:</strong> {esc(p)}</div>\n'
        elif 'concern' in p.lower() and len(p) < 80:
            pass  # titolo sezione, skip
        elif p.startswith('Observation:') or p.startswith('Impact:') or \
             p.startswith('Current Status:') or p.startswith('Concern:'):
            pass  # gestiti in AOB section

    # Contacts divisi per company
    contacts = d.get('contacts', [])
    iqs    = [c for c in contacts if 'SAIPEM' in c['company'].upper() or 'IQS' in c['company'].upper()]
    eew    = [c for c in contacts if 'EEW' in c['company'].upper()]
    sumi   = [c for c in contacts if 'SUMITOMO' in c['company'].upper()]
    qe     = [c for c in contacts if 'QE' in c['company'].upper() or 'QATAR' in c['company'].upper()]
    tpia   = [c for c in contacts if 'TPIA' in c['role'].upper()]

    def contact_rows(lst):
        html = ''
        for c in lst:
            initials = ''.join(w[0].upper() for w in c['name'].split() if w)[:2]
            email_part = f'<div style="font-size:10px;color:var(--text3)">{esc(c["email"])}</div>' if c['email'] else ''
            av_cls = 'avb' if 'SAIPEM' in c['company'].upper() or 'IQS' in c['company'].upper() \
                else 'avt' if 'EEW' in c['company'].upper() \
                else 'avp' if 'SUMITOMO' in c['company'].upper() \
                else 'ava'
            html += f'<div class="prow"><div class="av {av_cls}">{initials}</div>'
            html += f'<div class="pi"><div class="pn">{esc(c["name"])}</div>'
            html += f'<div class="pr2">{esc(c["role"])}</div>{email_part}</div></div>\n'
        return html

    # Activities table (Date | FR | ITP Pos. | Description | Pipe Nos | Result)
    acts_html = ''
    for a in d.get('activities', []):
        res = a.get('result', '')
        ok  = 'pg' if res.lower().startswith(('satisf','✔','pass','ok')) else 'pgr'
        acts_html += f'''<tr>
  <td class="mono">{esc(a["date"])}</td>
  <td class="mono">{esc(a.get("fr",""))}</td>
  <td><strong>{esc(a["itp"])}</strong></td>
  <td>{esc(a["desc"])}</td>
  <td style="font-size:11px">{esc(a.get("pipes",""))}</td>
  <td style="font-size:11px">{esc(res)}</td>
</tr>\n'''

    # Tally table
    tally_rows = ''
    for t in d.get('tally', []):
        tally_rows += f'''<tr>
  <td>{esc(t["date"])}</td><td class="mono">{esc(t["no"])}</td>
  <td>{esc(t["item"])}</td><td><strong>{esc(t["qty"])}</strong></td>
  <td>{esc(t["length"])}</td><td>{esc(t["saipem"])}</td>
  <td style="font-size:11px">{esc(t["remarks"])}</td>
</tr>\n'''
    tt = d.get('tally_total', {})
    if tt:
        tally_rows += f'''<tr class="rh">
  <td colspan="3">Total</td><td><strong>{esc(tt.get("qty",""))}</strong></td>
  <td>{esc(tt.get("length",""))}</td><td colspan="2"></td>
</tr>\n'''

    # Phases bars data for JS
    phases_js = '[\n'
    for p in d.get('phases', []):
        phases_js += f'  {{name:"{esc(p["name"])}",done:{p["completed"]},total:{p["total"]}}},\n'
    phases_js = phases_js.rstrip(',\n') + '\n]'

    # Yearly table
    yearly_rows = ''
    for row in d.get('yearly', []):
        cells = ''.join(f'<td>{esc(c)}</td>' for c in row)
        style = ' class="rh"' if row[0].upper() in ('TOTAL','TOTAL.') else ''
        yearly_rows += f'<tr{style}>{cells}</tr>\n'

    # Phase snapshot table
    phases_snap = ''
    for p in d.get('phases', []):
        phases_snap += f'''<tr>
  <td>{esc(p["name"])}</td><td>{esc(p["completed"])}</td>
  <td>{esc(p["remaining"])}</td><td>{esc(p["total"])}</td>
  <td><strong>{esc(p["pct"])}</strong></td>
</tr>\n'''

    # Dispatch PO table
    dispatch_po_rows = ''
    for po in d.get('dispatch_po', []):
        dispatch_po_rows += f'''<tr>
  <td class="mono">{esc(po["po"])}</td><td>{esc(po["released"])}</td>
  <td>{esc(po["length_mm"])}</td><td>{esc(po["dispatched"])}</td>
  <td><strong>{esc(po["ready"])}</strong></td>
</tr>\n'''
    dt2 = d.get('dispatch_total',{})
    dispatch_po_rows += f'''<tr class="rh">
  <td>Total</td><td>{esc(dt2.get("released","—"))}</td>
  <td>{esc(dt2.get("length_mm","—"))}</td><td>{esc(dt2.get("dispatched","—"))}</td>
  <td>{esc(dt2.get("ready","—"))}</td>
</tr>\n'''

    # Dispatch cats
    dispatch_cat_rows = ''
    for cat in d.get('dispatch_cats', []):
        dispatch_cat_rows += f'''<tr>
  <td>{esc(cat["cat"])}</td><td>{esc(cat["released"])}</td>
  <td>{esc(cat["length"])}</td><td>{esc(cat["dispatched"])}</td>
  <td>{esc(cat["ready"])}</td>
</tr>\n'''

    # Antwerp table
    antwerp_rows = ''
    for row in d.get('antwerp', []):
        style = ' class="rh"' if not re.match(r'^\d{1,2}$', row['sow']) and row['sow'] else ''
        antwerp_rows += f'<tr{style}><td>{esc(row["sow"])}</td><td>{esc(row["nos"])}</td><td>{esc(row["length"])}</td></tr>\n'

    # Materials
    mat_rows = ''
    for m2 in d.get('materials', []):
        bold = '<strong>' if m2['pct'].startswith('9') else ''
        ebold = '</strong>' if bold else ''
        mat_rows += f'<tr><td>{esc(m2["type"])}</td><td>{esc(m2["qty"])}</td><td>{bold}{esc(m2["pct"])}{ebold}</td></tr>\n'

    # Items summary
    items_rows = ''
    for it in d.get('items', []):
        s = it['status'].strip().upper()
        p_cls = 'pg' if s in ('COMPLETE','COMPLETED') else 'pa' if s == 'HALFWAY' else 'pb'
        p_label = 'Complete' if s in ('COMPLETE','COMPLETED') else it['status']
        comp = it['completion']
        comp_str = f"<strong>{comp}%</strong>" if s not in ('COMPLETE','COMPLETED') else f"{comp}%"
        items_rows += f'''<tr>
  <td class="mono">{esc(it["item"])}</td><td>{esc(it["pipes"])}</td>
  <td>{comp_str}</td><td><span class="pill {p_cls}">{p_label}</span></td>
</tr>\n'''

    # VAGB delivery
    vagb_rows = ''
    for row in d.get('vagb_delivery', []):
        bg = ' style="background:rgba(29,158,117,.07)"' if row['item'] == '00400' else ''
        vagb_rows += f'''<tr{bg}>
  <td class="mono">{esc(row["item"])}</td><td>{esc(row["total"])}</td>
  <td>{esc(row["delivered"])}</td><td>{esc(row["pending"])}</td>
  <td>{esc(row["ordered"])}</td><td><strong>{esc(row["pct"])}</strong></td>
  <td>{esc(row["first"])}</td><td>{esc(row["last"])}</td>
</tr>\n'''

    # VAGB schedule
    vagb_sched = ''
    for row in d.get('vagb_schedule', []):
        rcv = row.get('received','')
        p_cls = 'pg' if rcv and rcv not in ('-','') and int(rcv) > 0 else 'pgr'
        vagb_sched += f'''<tr>
  <td>W{esc(row["week"])}</td><td>{esc(row["dispatch"])}</td>
  <td>{esc(row["schedule"])}</td>
  <td>{('<span class="pill '+p_cls+'">'+esc(rcv)+'</span>') if rcv else '<span class="pill pgr">—</span>'}</td>
</tr>\n'''

    # HSE
    hse = d.get('hse', {})
    hse_rows = ''
    for k in ('1','2','3','4'):
        h = hse.get(k, {})
        desc = h.get('desc', '')
        yes  = h.get('yes', False)
        no   = h.get('no',  False)
        val = pill('Yes', 'pg') if yes else pill('No', 'pg') if no else '—'
        hse_rows += f'<div class="crow"><span class="clabel">{esc(k)}. {esc(desc)}</span>{val}</div>\n'
    hse_note = d.get('hse_note', 'No HSE issues this inspection visit/period. Nothing to report.')

    # AOB section
    aob_sections = {}
    current_section = None
    for p in aob:
        if re.match(r'\d+\.\s+', p):
            current_section = p
            aob_sections[current_section] = []
        elif current_section:
            aob_sections[current_section].append(p)

    aob_html = ''
    dot_map = {'Observation:':'dg','Area of concern:':'dr','Concern:':'da',
               'Impact:':'da','Current Status:':'db','Action Plan:':'dg',
               'Timeline:':'dg','MPFB':'dg','Line Pipes:':'da'}
    for section, paras in aob_sections.items():
        aob_html += f'<div style="font-size:11px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.05em;margin-bottom:8px">{esc(section)}</div>\n'
        aob_html += '<div style="font-size:12px;line-height:1.8;color:var(--text2);margin-bottom:14px">\n'
        for para in paras:
            d_cls = 'db'
            for key, cls in dot_map.items():
                if para.startswith(key): d_cls = cls; break
            aob_html += f'<div style="display:flex;gap:8px;margin-bottom:4px"><span class="dot {d_cls}" style="margin-top:7px"></span><span>{esc(para)}</span></div>\n'
        aob_html += '</div>\n<div class="sep"></div>\n'

    # Photos section HTML
    photos_section_html = ''
    for step, indices in photo_groups.items():
        group_captions = photos[indices[0]]['caption'] if indices else ''
        # Titolo sezione
        step_title = f'ITP Step {step} — '
        # Deduce titolo dal caption
        if indices:
            cap = photos[indices[0]]['caption']
            m2 = re.search(r'ITP Step \d+ — (.+)', cap, re.IGNORECASE)
            step_title = f'ITP Step {step} — {m2.group(1).split("—")[0].strip()}' if m2 else f'ITP Step {step}'
        grid_id = f'pg-step{step}'
        photos_section_html += f'''<div class="card">
  <div class="ct"><i class="ti ti-photo"></i>{esc(step_title)}</div>
  <div class="photo-grid" id="{grid_id}"></div>
</div>\n'''

    # Photo render JS
    photo_render_js = ''
    for step, indices in photo_groups.items():
        grid_id = f'pg-step{step}'
        photo_render_js += f'  const g{step}=document.getElementById("{grid_id}");\n'
        photo_render_js += f'  [{",".join(str(i) for i in indices)}].forEach(i=>g{step}&&g{step}.appendChild(mkCard(PHOTOS[i][0],PHOTOS[i][1],i)));\n'

    # Enclosures
    enc_html = ''
    for e in d.get('enclosures', []):
        enc_html += f'''<div class="crow">
  <span style="font-family:var(--mono);font-size:10px;color:var(--text3);min-width:20px">{esc(e["num"])}</span>
  <span class="clabel">{esc(e["title"])}</span>
  <span class="pill pgr">{esc(e["type"])}</span>
</div>\n'''

    # Completion forecast (tabella 15)
    def fc_pill(s):
        s = (s or '').upper()
        if 'PROGRESS' in s: return '<span class="pill pa">In progress</span>'
        if 'COMPLET' in s or 'DONE' in s: return '<span class="pill pg">Completed</span>'
        if 'DELAY' in s or 'LATE' in s: return '<span class="pill pr">Delayed</span>'
        return f'<span class="pill pgr">{esc(s.title())}</span>'
    forecast_rows = ''
    for f in d.get('forecast', []):
        forecast_rows += f'''<tr>
  <td class="mono">{esc(f["phase"])}</td>
  <td>{esc(f["actual"])}</td><td>{esc(f["target"])}</td><td>{esc(f["remaining"])}</td>
  <td>{esc(f["rate"])}</td><td>{esc(f["prod7"])}</td>
  <td><strong>{esc(f["est"])}</strong></td><td>{esc(f["po_date"])}</td>
  <td>{fc_pill(f["status"])}</td>
</tr>\n'''

    # Confronto data PO (tabella 17)
    po_cmp_rows = ''
    for label, val in d.get('po_comparison', []):
        u = (label+val).upper()
        cls = 'pr' if ('DELAY' in u or '+' in val) else 'pg' if ('AHEAD' in u or 'ON TRACK' in u) else 'pgr'
        is_status = 'STATUS' in label.upper() or 'DEVIATION' in label.upper()
        valhtml = f'<span class="pill {cls}">{esc(val)}</span>' if is_status else f'<strong>{esc(val)}</strong>'
        po_cmp_rows += f'<div class="crow"><span class="clabel">{esc(label)}</span>{valhtml}</div>\n'

    # Produzione per periodo (tabelle 11/12)
    def prod_table_rows(rows, ncols=9):
        out = ''
        for r in rows[1:]:
            cells = (r + ['']*ncols)[:ncols]
            tds = ''.join(f'<td>{esc(c)}</td>' for c in cells)
            out += f'<tr>{tds}</tr>\n'
        return out
    prod_week_rows  = prod_table_rows(d.get('prod_week', []))
    prod_month_rows = prod_table_rows(d.get('prod_month', []))

    # Prove di laboratorio / meccaniche (tabella 20)
    def test_pill(s):
        s = str(s)
        if 'PASS' in s.upper() or '✔' in s: return f'<span class="pill pg">{esc(s)}</span>'
        if 'FAIL' in s.upper() or '✘' in s: return f'<span class="pill pr">{esc(s)}</span>'
        if 'PEND' in s.upper(): return f'<span class="pill pa">{esc(s)}</span>'
        return f'<span class="pill pgr">{esc(s)}</span>'
    mech = d.get('mech', [])
    mech_rows = ''
    for r in mech[1:]:
        c = (r + ['']*9)[:9]
        mech_rows += f'''<tr>
  <td class="mono">{esc(c[0])}</td><td>{esc(c[1])}</td><td class="mono">{esc(c[2])}</td>
  <td class="mono">{esc(c[3])}</td><td class="mono">{esc(c[4])}</td><td class="mono">{esc(c[5])}</td>
  <td class="mono">{esc(c[6])}</td><td>{esc(c[7])}</td><td>{test_pill(c[8])}</td>
</tr>\n'''
    mech_pass = sum(1 for r in mech[1:] if 'PASS' in ' '.join(r).upper() or '✔' in ' '.join(r))
    mech_total = max(len(mech)-1, 0)

    # Generica: header + righe da una tabella grezza
    def raw_table(rows):
        if not rows: return '', ''
        th = ''.join(f'<th>{esc(c)}</th>' for c in rows[0])
        ncol = len(rows[0])
        body = ''
        for r in rows[1:]:
            c = (r + ['']*ncol)[:ncol]
            body += '<tr>' + ''.join(f'<td>{esc(x)}</td>' for x in c) + '</tr>\n'
        return th, body
    corr_th, corr_body   = raw_table(d.get('mech_corrosion', []))
    dwtt_th, dwtt_body   = raw_table(d.get('mech_dwtt', []))
    tens_th, tens_body   = raw_table(d.get('mech_tensile', []))

    # MDR / certificati (tabella 29) — presente solo in alcune settimane
    has_mdr = bool(d.get('mdr_pct'))
    mdr_pct     = d.get('mdr_pct', '—')
    mdr_missing = d.get('mdr_missing', '—')
    mdr_rows = ''
    for m in d.get('mdr_detail', []):
        mdr_rows += f'''<tr>
  <td>{esc(m["status"])}</td><td>{esc(m["qty"])}</td>
  <td><strong>{esc(m["pct"])}</strong></td><td>{esc(m["op"])}</td>
</tr>\n'''
    # Frammenti condizionali MDR
    mdr_kpi_cards = (f'''<div class="kpi kg"><div class="kl">MDR completion</div><div class="kv">{mdr_pct}</div><div class="ks">signed certificates</div></div>
    <div class="kpi kr"><div class="kl">Missing certificates</div><div class="kv">{mdr_missing}</div><div class="ks">to complete MDR</div></div>''' if has_mdr else '')
    mdr_card = (f'''<div class="card">
    <div class="ct"><i class="ti ti-certificate"></i>MDR — certificate status</div>
    <div class="tw"><table><thead><tr><th>Plate / certificate status</th><th>Qty (plates)</th><th>%</th><th>Operational status</th></tr></thead>
    <tbody>{mdr_rows}</tbody></table></div>
  </div>''' if has_mdr else '')
    # Card MDR per Overview (fallback a VAGB delta se assente)
    overview_mdr_card = (f'<div class="kpi kg"><div class="kl">MDR certificates</div><div class="kv">{mdr_pct}</div><div class="ks">{mdr_missing} missing</div></div>'
                         if has_mdr else
                         f'<div class="kpi kb"><div class="kl">VAGB delta</div><div class="kv">{vagb_delta}</div><div class="ks">vs 100% ordered</div></div>')

    # NOI prossima visita (tabella 27)
    noi_rows = ''
    for x in d.get('noi', []):
        noi_rows += f'''<tr>
  <td class="mono">{esc(x["noi"])}</td><td>{esc(x["activity"])}</td>
  <td class="mono">{esc(x["date"])}</td><td>{esc(x["time"])}</td>
  <td style="font-size:11px">{esc(x["loc"])}</td>
  <td>{status_pill(x["status"])}</td>
</tr>\n'''

    # Documenti ITP usati (tabella 4)
    docs_rows = ''
    for e in d.get('docs', []):
        if 'section' in e:
            docs_rows += f'<tr class="rh"><td colspan="4">{esc(e["section"])}</td></tr>\n'
        else:
            docs_rows += f'''<tr>
  <td class="mono">{esc(e["num"])}</td><td>{esc(e["title"])}</td>
  <td>{esc(e["rev"])}</td><td>{status_pill(e["status"])}</td>
</tr>\n'''

    # Inspected pipes per PO (tabella 9)
    inspected_length = d.get('inspected_length', '—')
    inspected_rows = ''
    for r in d.get('inspected_po', []):
        inspected_rows += f'<tr><td class="mono">{esc(r["po"])}</td><td>{esc(r["nos"])}</td><td>{esc(r["length"])}</td></tr>\n'
    if d.get('inspected_total'):
        inspected_rows += f'<tr class="rh"><td>Total</td><td><strong>{esc(d.get("inspected_total","—"))}</strong></td><td>{esc(inspected_length)}</td></tr>\n'

    # VAGB plates KPI table (tabella 28)
    vagb_kpi_rows = ''
    for k, v in d.get('vagb_kpi', {}).items():
        vagb_kpi_rows += f'''<tr>
  <td>{esc(k)}</td><td><strong>{esc(v.get("value",""))}</strong></td>
  <td>{esc(v.get("target",""))}</td><td>{esc(v.get("delta",""))}</td>
</tr>\n'''

    # Valori sintetici per KPI/overview dal forecast e confronto PO
    fc_est = d['forecast'][-1]['est'] if d.get('forecast') else '—'   # Final Inspection
    po_dev = next((v for l,v in d.get('po_comparison',[]) if 'DEVIATION' in l.upper()), '—')
    po_status = next((v for l,v in d.get('po_comparison',[]) if l.upper().strip()=='STATUS'), '')

    # ── CSS (identico al template pulito) ────────────────────────────────────
    css = """
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --accent:#1D9E75;--accent2:#378ADD;--warn:#EF9F27;--danger:#E24B4A;
  --bg:#ffffff;--bg2:#f5f5f3;--bg3:#eeede8;
  --text:#1a1a18;--text2:#5f5e5a;--text3:#888780;
  --border:rgba(0,0,0,0.09);--border2:rgba(0,0,0,0.17);
  --r:8px;--rl:12px;
  --font:-apple-system,'Segoe UI',system-ui,sans-serif;
  --mono:'SF Mono','Consolas','Liberation Mono',monospace;
  --gn-bg:#EAF3DE;--gn-t:#27500A;--am-bg:#FAEEDA;--am-t:#633806;
  --rd-bg:#FCEBEB;--rd-t:#A32D2D;--bl-bg:#E6F1FB;--bl-t:#0C447C;
  --tl-bg:#9FE1CB;--tl-t:#085041;--pu-bg:#CECBF6;--pu-t:#3C3489;
  --sw:210px;
}
@media(prefers-color-scheme:dark){
  :root{--bg:#1c1c1a;--bg2:#252523;--bg3:#2e2e2b;--text:#f0efe8;--text2:#b4b2a9;--text3:#888780;
    --border:rgba(255,255,255,0.09);--border2:rgba(255,255,255,0.17);
    --gn-bg:#173404;--gn-t:#C0DD97;--am-bg:#412402;--am-t:#FAC775;
    --rd-bg:#501313;--rd-t:#F7C1C1;--bl-bg:#042C53;--bl-t:#B5D4F4;
    --tl-bg:#04342C;--tl-t:#9FE1CB;--pu-bg:#26215C;--pu-t:#CECBF6;}
}
html,body{height:100%;font-family:var(--font);font-size:13px;background:var(--bg3);color:var(--text);line-height:1.5}
.sidebar{position:fixed;top:0;left:0;bottom:0;width:var(--sw);background:var(--bg2);border-right:0.5px solid var(--border2);display:flex;flex-direction:column;z-index:100;transition:transform .28s cubic-bezier(.4,0,.2,1)}
.sb-backdrop{display:none;position:fixed;inset:0;background:rgba(0,0,0,.5);z-index:99}
.sb-backdrop.open{display:block}
.sb-logo{padding:18px 16px 14px;border-bottom:0.5px solid var(--border)}
.sb-tag{font-size:10px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.07em;margin-bottom:3px}
.sb-title{font-size:14px;font-weight:500;color:var(--text)}
.sb-sub{font-size:11px;color:var(--text3);margin-top:2px}
.sb-week{margin:10px 12px;background:var(--bg3);border-radius:var(--r);padding:9px 11px;border:0.5px solid var(--border)}
.sb-wl{font-size:10px;color:var(--text3);text-transform:uppercase;letter-spacing:.05em}
.sb-wv{font-size:12px;font-weight:500;color:var(--text);margin-top:1px}
.sb-nav{flex:1;padding:8px 0;overflow-y:auto}
.sb-sec{font-size:10px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.07em;padding:12px 16px 4px}
.sb-btn{display:flex;align-items:center;gap:8px;width:100%;padding:7px 16px;border:none;background:transparent;font-family:var(--font);font-size:12px;color:var(--text2);cursor:pointer;text-align:left;border-left:2px solid transparent;transition:all .14s}
.sb-btn i{font-size:15px;opacity:.6}
.sb-btn:hover{background:var(--bg);color:var(--text)}
.sb-btn.active{background:var(--bg);color:var(--text);border-left-color:var(--accent);font-weight:500}
.sb-btn.active i{opacity:1;color:var(--accent)}
.sb-foot{padding:14px 16px;border-top:0.5px solid var(--border)}
.sb-irlabel{font-size:10px;color:var(--text3);text-transform:uppercase;letter-spacing:.05em}
.sb-irn{font-size:12px;color:var(--text2);margin-top:2px;font-family:var(--mono)}
.main{margin-left:var(--sw);height:100vh;display:flex;flex-direction:column}
.topbar{padding:11px 24px;background:var(--bg);border-bottom:0.5px solid var(--border2);display:flex;align-items:center;justify-content:space-between;gap:8px;flex-shrink:0}
.tb-title{font-size:14px;font-weight:500;color:var(--text)}
.tb-badges{display:flex;gap:6px}
.badge{font-size:10px;font-weight:500;padding:3px 9px;border-radius:100px;white-space:nowrap}
.bw{background:var(--am-bg);color:var(--am-t)}.bg{background:var(--gn-bg);color:var(--gn-t)}.bi{background:var(--bl-bg);color:var(--bl-t)}
.hamburger{display:none;align-items:center;justify-content:center;width:34px;height:34px;border:none;background:var(--bg2);border-radius:var(--r);cursor:pointer;border:0.5px solid var(--border)}
.hamburger i{font-size:18px;color:var(--text2)}
.content{padding:20px 24px;overflow-y:auto;flex:1;min-height:0}
.section{display:none}.section.active{display:block}
.kpi-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-bottom:14px}
.kpi{background:var(--bg);border:0.5px solid var(--border);border-radius:var(--r);padding:12px 14px;position:relative}
.kpi::before{content:'';position:absolute;top:0;left:0;width:3px;height:100%}
.kg::before{background:var(--accent)}.kb::before{background:var(--accent2)}.kw::before{background:var(--warn)}.kr::before{background:var(--danger)}
.kl{font-size:10px;color:var(--text3);text-transform:uppercase;letter-spacing:.04em;margin-bottom:5px}
.kv{font-size:22px;font-weight:500;color:var(--text);line-height:1;letter-spacing:-.02em}
.ks{font-size:10px;color:var(--text3);margin-top:3px}
.alert{padding:10px 13px;border-left:2px solid;margin-bottom:10px;font-size:12px;line-height:1.6;border-radius:0 var(--r) var(--r) 0}
.aw{background:var(--am-bg);border-color:var(--warn);color:var(--am-t)}
.ai{background:var(--bl-bg);border-color:var(--accent2);color:var(--bl-t)}
.ag{background:var(--gn-bg);border-color:var(--accent);color:var(--gn-t)}
.card{background:var(--bg);border:0.5px solid var(--border);border-radius:var(--rl);padding:16px 18px;margin-bottom:12px}
.ct{font-size:10px;font-weight:500;color:var(--text3);text-transform:uppercase;letter-spacing:.07em;margin-bottom:13px;display:flex;align-items:center;gap:7px}
.ct i{font-size:14px;opacity:.55}
.two{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:12px}
.tw{overflow-x:auto}
table{width:100%;border-collapse:collapse;font-size:12px}
th{font-size:10px;font-weight:500;color:var(--text3);text-align:left;padding:7px 10px;border-bottom:0.5px solid var(--border2);background:var(--bg2);text-transform:uppercase;letter-spacing:.04em;white-space:nowrap}
td{padding:8px 10px;border-bottom:0.5px solid var(--border);color:var(--text);vertical-align:middle}
tr:last-child td{border-bottom:none}
tbody tr:hover td{background:var(--bg2)}
.mono{font-family:var(--mono);font-size:11px;color:var(--text2)}
.rh td{background:var(--bg2);font-weight:500}
.pill{display:inline-block;font-size:10px;font-weight:500;padding:2px 8px;border-radius:100px;white-space:nowrap}
.pg{background:var(--gn-bg);color:var(--gn-t)}.pa{background:var(--am-bg);color:var(--am-t)}
.pr{background:var(--rd-bg);color:var(--rd-t)}.pb{background:var(--bl-bg);color:var(--bl-t)}
.pt{background:var(--tl-bg);color:var(--tl-t)}.pp{background:var(--pu-bg);color:var(--pu-t)}
.pgr{background:var(--bg3);color:var(--text2)}
.ph-row{display:flex;align-items:center;gap:10px;margin-bottom:10px}
.ph-row:last-child{margin-bottom:0}
.ph-name{font-size:12px;color:var(--text2);min-width:130px}
.ph-track{flex:1;height:6px;background:var(--bg3);border-radius:3px;overflow:hidden;border:0.5px solid var(--border)}
.ph-bar{height:100%;border-radius:3px;background:var(--accent);transition:width .9s cubic-bezier(.22,1,.36,1)}
.ph-pct{font-size:11px;font-weight:500;color:var(--text);min-width:38px;text-align:right}
.ph-cnt{font-size:10px;color:var(--text3);min-width:70px;text-align:right}
.prow{display:flex;align-items:center;gap:10px;padding:8px 0;border-bottom:0.5px solid var(--border)}
.prow:last-child{border-bottom:none}
.av{width:32px;height:32px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:10px;font-weight:500;flex-shrink:0}
.avb{background:var(--bl-bg);color:var(--bl-t)}.avt{background:var(--tl-bg);color:var(--tl-t)}
.avp{background:var(--pu-bg);color:var(--pu-t)}.ava{background:var(--am-bg);color:var(--am-t)}
.avg{background:var(--bg3);color:var(--text2)}
.pi{flex:1;min-width:0}
.pn{font-size:12px;font-weight:500;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.pr2{font-size:10px;color:var(--text2);margin-top:1px}
.dot{display:inline-block;width:6px;height:6px;border-radius:50%;flex-shrink:0;margin-right:6px}
.dg{background:var(--accent)}.db{background:var(--accent2)}.da{background:var(--warn)}.dr{background:var(--danger)}
.irn-list{max-height:300px;overflow-y:auto}
.irn-row{display:flex;align-items:center;gap:8px;padding:5px 0;border-bottom:0.5px solid var(--border);font-size:12px}
.irn-row:last-child{border-bottom:none}
.irn-id{font-family:var(--mono);font-size:10px;color:var(--text3);min-width:130px}
.irn-desc{color:var(--text);flex:1}
.crow{display:flex;align-items:center;gap:10px;padding:7px 0;border-bottom:0.5px solid var(--border);font-size:12px}
.crow:last-child{border-bottom:none}
.clabel{flex:1;color:var(--text2)}
.sep{height:0.5px;background:var(--border);margin:12px 0}
.arow{display:flex;align-items:center;gap:12px;padding:8px 0;border-bottom:0.5px solid var(--border)}
.arow:last-child{border-bottom:none}
.adate{font-family:var(--mono);font-size:11px;color:var(--text3);min-width:80px}
.aitp{font-size:10px;font-weight:500;color:var(--text3);min-width:36px;text-align:center}
.adesc{flex:1;font-size:12px;color:var(--text)}
.anote{font-size:11px;color:var(--text2)}
.sdot{width:8px;height:8px;border-radius:50%;flex-shrink:0;background:var(--warn)}
.photo-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:8px;margin-bottom:12px}
.photo-card{background:var(--bg2);border:0.5px solid var(--border);border-radius:var(--r);overflow:hidden}
.photo-img{width:100%;height:130px;object-fit:cover;display:block;cursor:zoom-in;transition:opacity .15s}
.photo-img:hover{opacity:.88}
.photo-cap{padding:7px 10px;font-size:10px;color:var(--text2)}
.lb{display:none;position:fixed;inset:0;background:rgba(0,0,0,.88);z-index:9999;align-items:center;justify-content:center;flex-direction:column;gap:12px}
.lb.open{display:flex}
.lb img{max-width:90vw;max-height:82vh;border-radius:4px;object-fit:contain}
.lb-cap{color:#fff;font-size:12px;opacity:.8}
.lb-close{position:absolute;top:16px;right:20px;background:none;border:none;color:#fff;font-size:28px;cursor:pointer;opacity:.7}
.lb-nav{position:absolute;top:50%;transform:translateY(-50%);background:rgba(255,255,255,.15);border:none;color:#fff;font-size:22px;cursor:pointer;padding:10px 14px;border-radius:4px}
#lb-prev{left:16px}#lb-next{right:16px}
@media(max-width:768px){
  .sidebar{transform:translateX(-100%);width:80vw;max-width:260px;box-shadow:4px 0 32px rgba(0,0,0,.25)}
  .sidebar.open{transform:translateX(0)}
  .main{margin-left:0}
  .topbar{padding:10px 14px}
  .tb-badges{display:none}
  .hamburger{display:flex}
  .content{padding:12px 14px}
  .kpi-grid{grid-template-columns:repeat(2,1fr);gap:8px}
  .two{grid-template-columns:1fr}
  .card{padding:12px 14px}
  table{min-width:460px;font-size:11px}
  th{font-size:10px;padding:6px 8px}td{padding:6px 8px;font-size:11px}
  .ph-name{min-width:100px;font-size:11px}.ph-cnt{display:none}
  .photo-grid{grid-template-columns:repeat(2,1fr)}
  .photo-img{height:110px}
}
@media(max-width:480px){
  .content{padding:10px}.kpi-grid{gap:6px}.kpi{padding:9px 10px}.kv{font-size:18px}
  .card{padding:10px 12px}table{min-width:380px}th,td{padding:5px 7px}
  .photo-img{height:90px}
}
@media print{
  .sidebar,.sb-backdrop,.hamburger{display:none!important}
  .main{margin-left:0;height:auto}
  .content{overflow:visible;padding:1rem}
  .section{display:block!important;page-break-before:always}
  .section:first-of-type{page-break-before:auto}
}
""" + LOGIN_GATE_CSS

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>IR-{ir:03d}-EEW COMP3B — Weekly Report {slash_end}</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/@tabler/icons-webfont@3.19.0/dist/tabler-icons.min.css">
<style>{css}</style>
</head>
<body>

{LOGIN_GATE_HTML}

<div class="sb-backdrop" id="sb-backdrop"></div>
<aside class="sidebar" id="sidebar">
  <div class="sb-logo">
    <div class="sb-tag">NFPS COMP3B</div>
    <div class="sb-title">EEW Weekly Report</div>
    <div class="sb-sub">Erndtebrück · Saipem / QE LNG</div>
  </div>
  <div class="sb-week">
    <div class="sb-wl">Reference week</div>
    <div class="sb-wv">{sidebar_w}</div>
  </div>
  <nav class="sb-nav">
    <div class="sb-sec">Report</div>
    <button class="sb-btn active" onclick="nav('overview',this)"><i class="ti ti-layout-dashboard"></i>Overview</button>
    <button class="sb-btn" onclick="nav('production',this)"><i class="ti ti-chart-bar"></i>Production &amp; forecast</button>
    <button class="sb-btn" onclick="nav('dispatch',this)"><i class="ti ti-truck"></i>Dispatch</button>
    <button class="sb-btn" onclick="nav('activities',this)"><i class="ti ti-calendar-event"></i>Weekly activities</button>
    <button class="sb-btn" onclick="nav('labtests',this)"><i class="ti ti-flask"></i>Lab &amp; tests</button>
    <div class="sb-sec">Supplier</div>
    <button class="sb-btn" onclick="nav('plates',this)"><i class="ti ti-layers"></i>Plates &amp; MDR</button>
    <button class="sb-btn" onclick="nav('nextvisit',this)"><i class="ti ti-calendar-due"></i>Next visit</button>
    <div class="sb-sec">Admin</div>
    <button class="sb-btn" onclick="nav('docs',this)"><i class="ti ti-folder"></i>ITP documents</button>
    <button class="sb-btn" onclick="nav('photos',this)"><i class="ti ti-photo"></i>Photos</button>
    <button class="sb-btn" onclick="nav('hse',this)"><i class="ti ti-shield"></i>HSE / AOB</button>
    <button class="sb-btn" onclick="nav('contacts',this)"><i class="ti ti-users"></i>Contacts</button>
  </nav>
  <div class="sb-foot">
    <div class="sb-irlabel">Inspection report</div>
    <div class="sb-irn">IR-{ir:03d}-EEW COMP3B</div>
  </div>
</aside>

<div class="main">
  <div class="topbar">
    <button class="hamburger" id="hb-btn"><i class="ti ti-menu-2"></i></button>
    <span class="tb-title" id="tb-title">Overview</span>
    <div class="tb-badges">
      <span class="badge bw">&#9651; Production slowed</span>
      <span class="badge bg">NCR: all closed</span>
      <span class="badge bi">Survey (S)</span>
      <span class="badge admin-visible admin-badge">ADMIN</span>
      <button class="badge admin-visible admin-badge" id="admin-print-btn" onclick="window.print()" style="cursor:pointer">&#128424; Stampa / PDF</button>
    </div>
  </div>
  <div class="content">

<!-- OVERVIEW -->
<div id="overview" class="section active">
  <div class="kpi-grid">
    <div class="kpi kb"><div class="kl">VAGB plates at EEW</div><div class="kv">{vagb_pct}</div><div class="ks">{vagb_nos} / {vagb_ordered} pcs</div></div>
    {overview_mdr_card}
    <div class="kpi kg"><div class="kl">Pipes released</div><div class="kv">{pipes_released}</div><div class="ks">3 purchase orders</div></div>
    <div class="kpi kg"><div class="kl">Ready for dispatch</div><div class="kv">{pipes_ready}</div><div class="ks">awaiting shipment</div></div>
  </div>
  <div class="kpi-grid">
    <div class="kpi kb"><div class="kl">Dispatched</div><div class="kv">{pipes_dispatched}</div><div class="ks">to final destination</div></div>
    <div class="kpi kg"><div class="kl">Pipes inspected</div><div class="kv">{pipes_inspected}</div><div class="ks">total</div></div>
    <div class="kpi kw"><div class="kl">Est. completion</div><div class="kv">{fc_est}</div><div class="ks">Item 00400 — Final Insp.</div></div>
    <div class="kpi kr"><div class="kl">Deviation vs PO</div><div class="kv">{po_dev}</div><div class="ks">{esc(po_status)}</div></div>
  </div>
  {alerts_html}
  <div class="two">
    <div class="card">
      <div class="ct"><i class="ti ti-target"></i>Comparison with PO completion date</div>
      {po_cmp_rows}
    </div>
    <div class="card">
      <div class="ct"><i class="ti ti-truck-delivery"></i>Dispatch summary — as of {slash_end}</div>
      <div class="tw"><table><thead><tr><th>Saipem PO</th><th>Released</th><th>Length (m)</th><th>Dispatched</th><th>Ready</th></tr></thead>
      <tbody>{dispatch_po_rows}</tbody></table></div>
    </div>
  </div>
</div>

<!-- PRODUCTION & FORECAST -->
<div id="production" class="section">
  <div class="card">
    <div class="ct"><i class="ti ti-target-arrow"></i>Completion forecast by phase — Item 00400</div>
    <div class="tw"><table><thead><tr>
      <th>Phase</th><th>Actual</th><th>Target</th><th>Remaining</th><th>Rate (pcs/d)</th>
      <th>Produced (7d)</th><th>Est. completion</th><th>PO date</th><th>Status</th>
    </tr></thead><tbody>{forecast_rows}</tbody></table></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-calendar-stats"></i>Production by year</div>
    <div class="tw"><table><thead><tr><th>Year</th><th>Milling</th><th>Weld. Base</th><th>Weld. Clad</th><th>Hydro</th><th>UT</th><th>RT</th><th>PT</th><th>Final Insp.</th></tr></thead>
    <tbody>{yearly_rows}</tbody></table></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-calendar-month"></i>Production by month</div>
    <div class="tw"><table><thead><tr><th>Month</th><th>Milling</th><th>Weld. Base</th><th>Weld. Clad</th><th>Hydro</th><th>UT</th><th>RT</th><th>PT</th><th>Final Insp.</th></tr></thead>
    <tbody>{prod_month_rows}</tbody></table></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-calendar-week"></i>Production by week</div>
    <div class="tw" style="max-height:360px;overflow-y:auto"><table><thead><tr><th>Week</th><th>Milling</th><th>Weld. Base</th><th>Weld. Clad</th><th>Hydro</th><th>UT</th><th>RT</th><th>PT</th><th>Final Insp.</th></tr></thead>
    <tbody>{prod_week_rows}</tbody></table></div>
  </div>
</div>

<!-- DISPATCH -->
<div id="dispatch" class="section">
  <div class="kpi-grid">
    <div class="kpi kg"><div class="kl">Total released</div><div class="kv">{pipes_released}</div><div class="ks">3 POs</div></div>
    <div class="kpi kb"><div class="kl">Dispatched</div><div class="kv">{pipes_dispatched}</div><div class="ks">to destination</div></div>
    <div class="kpi kw"><div class="kl">Ready for dispatch</div><div class="kv">{pipes_ready}</div><div class="ks">awaiting</div></div>
    <div class="kpi kg"><div class="kl">Pipes inspected</div><div class="kv">{pipes_inspected}</div><div class="ks">{inspected_length} m</div></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-truck"></i>Summary by PO — as of {slash_end}</div>
    <div class="tw"><table><thead><tr><th>Saipem PO</th><th>Released</th><th>Length (m)</th><th>Dispatched</th><th>Ready</th></tr></thead>
    <tbody>{dispatch_po_rows}</tbody></table></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-checkup-list"></i>Inspected pipes by PO</div>
    <div class="tw"><table><thead><tr><th>Saipem PO</th><th>Inspected (Nos)</th><th>Length (m)</th></tr></thead>
    <tbody>{inspected_rows}</tbody></table></div>
  </div>
</div>

<!-- WEEKLY ACTIVITIES -->
<div id="activities" class="section">
  <div class="card">
    <div class="ct"><i class="ti ti-calendar-event"></i>Activities — {range_en}</div>
    <div style="font-size:11px;color:var(--text3);margin-bottom:12px">ITP: 5129_20-4600015016-00054 Rev.06 AFC</div>
    <div class="tw"><table><thead><tr>
      <th>Date</th><th>FR No.</th><th>ITP Pos.</th><th>Activity description</th><th>Pipe Nos / Sample</th><th>Result / Remarks</th>
    </tr></thead><tbody>{acts_html}</tbody></table></div>
  </div>
  {f'''<div class="card">
    <div class="ct"><i class="ti ti-clipboard-list"></i>Tally Lists — Weekly Summary</div>
    <div class="tw"><table><thead><tr>
      <th>Date</th><th>Tally List No.</th><th>Item</th><th>Qty (pipes)</th>
      <th>Length (mm)</th><th>Saipem / QE Item</th><th>Remarks</th>
    </tr></thead><tbody>{tally_rows}</tbody></table></div>
  </div>''' if tally_rows else ''}
</div>

<!-- LAB & TESTS -->
<div id="labtests" class="section">
  <div class="kpi-grid">
    <div class="kpi kg"><div class="kl">Tests this week</div><div class="kv">{mech_total}</div><div class="ks">lab &amp; mechanical</div></div>
    <div class="kpi kg"><div class="kl">Passed</div><div class="kv">{mech_pass}</div><div class="ks">within criteria</div></div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-flask"></i>Lab &amp; mechanical tests — {range_en}</div>
    <div class="tw"><table><thead><tr>
      <th>Date</th><th>Location</th><th>Pipe No.</th><th>Plate No.</th><th>CS Heat</th><th>CRA Heat</th><th>Sample</th><th>Test type</th><th>Result</th>
    </tr></thead><tbody>{mech_rows}</tbody></table></div>
  </div>
  {f'''<div class="card">
    <div class="ct"><i class="ti ti-droplet"></i>Corrosion test detail (ASTM G48 / G28)</div>
    <div class="tw"><table><thead><tr>{corr_th}</tr></thead><tbody>{corr_body}</tbody></table></div>
  </div>''' if corr_body else ''}
  {f'''<div class="card">
    <div class="ct"><i class="ti ti-temperature"></i>DWTT / shear area</div>
    <div class="tw"><table><thead><tr>{dwtt_th}</tr></thead><tbody>{dwtt_body}</tbody></table></div>
  </div>''' if dwtt_body else ''}
  {f'''<div class="card">
    <div class="ct"><i class="ti ti-ruler-2"></i>Tensile / hardness</div>
    <div class="tw"><table><thead><tr>{tens_th}</tr></thead><tbody>{tens_body}</tbody></table></div>
  </div>''' if tens_body else ''}
</div>

<!-- PLATES & MDR -->
<div id="plates" class="section">
  <div class="kpi-grid">
    <div class="kpi kb"><div class="kl">Plates at EEW</div><div class="kv">{vagb_nos}</div><div class="ks">of {vagb_ordered} ordered</div></div>
    <div class="kpi kg"><div class="kl">Delivery rate</div><div class="kv">{vagb_pct}</div><div class="ks">delta {vagb_delta}</div></div>
    {mdr_kpi_cards}
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-layers"></i>VAGB plates — KPI</div>
    <table><thead><tr><th>KPI / Indicator</th><th>Current</th><th>Target</th><th>Delta</th></tr></thead>
    <tbody>{vagb_kpi_rows}</tbody></table>
  </div>
  {mdr_card}
</div>

<!-- NEXT VISIT -->
<div id="nextvisit" class="section">
  <div class="card">
    <div class="ct"><i class="ti ti-calendar-due"></i>Planned inspections — next visit</div>
    <div class="tw"><table><thead><tr>
      <th>NOI N°</th><th>Activity</th><th>Date</th><th>Time</th><th>Location</th><th>Status</th>
    </tr></thead><tbody>{noi_rows}</tbody></table></div>
  </div>
</div>

<!-- ITP DOCUMENTS -->
<div id="docs" class="section">
  <div class="card">
    <div class="ct"><i class="ti ti-folder"></i>Used documents for inspection — {len(d.get("docs", []))} entries</div>
    <div class="tw" style="max-height:560px;overflow-y:auto"><table><thead><tr><th>Doc. No.</th><th>Title — description</th><th>Rev.</th><th>Status</th></tr></thead>
    <tbody>{docs_rows}</tbody></table></div>
  </div>
</div>

<!-- PHOTOS -->
<div id="photos" class="section">
  <div class="alert ai">Photography inside test houses is strictly forbidden.</div>
  {photos_section_html}
  <div class="card">
    <div class="ct"><i class="ti ti-paperclip"></i>IR enclosures</div>
    {enc_html}
  </div>
</div>

<!-- HSE / AOB -->
<div id="hse" class="section">
  <div class="card">
    <div class="ct"><i class="ti ti-shield"></i>HSE checklist — {range_en}</div>
    {hse_rows}
    <div style="margin-top:10px" class="alert ag">{esc(hse_note)}</div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-alert-triangle"></i>Areas of concern</div>
    {aob_html}
  </div>
</div>

<!-- CONTACTS -->
<div id="contacts" class="section">
  <div class="two">
    <div class="card">
      <div class="ct"><i class="ti ti-users"></i>IQS / Saipem</div>
      {contact_rows(iqs)}
    </div>
    <div class="card">
      <div class="ct"><i class="ti ti-building"></i>Sumitomo / QE LNG / TPIA</div>
      {contact_rows(sumi + qe + tpia)}
    </div>
  </div>
  <div class="card">
    <div class="ct"><i class="ti ti-building-factory"></i>EEW Pipe Production Erndtebrück</div>
    <div style="display:grid;grid-template-columns:1fr 1fr">
      <div>{contact_rows(eew[:len(eew)//2+1])}</div>
      <div>{contact_rows(eew[len(eew)//2+1:])}</div>
    </div>
  </div>
</div>

  </div>
</div>

<div id="lb" class="lb">
  <button class="lb-close" id="lb-close">&times;</button>
  <button class="lb-nav" id="lb-prev">&#8592;</button>
  <img id="lb-img" src="" alt="">
  <div id="lb-cap" class="lb-cap"></div>
  <button class="lb-nav" id="lb-next">&#8594;</button>
</div>

<script>
const TITLES={{overview:'Overview',production:'Production & forecast',dispatch:'Dispatch',
  activities:'Weekly activities',labtests:'Lab & tests',plates:'Plates & MDR',
  nextvisit:'Next visit',docs:'ITP documents',photos:'Photos',hse:'HSE / AOB',contacts:'Contacts'}};

function nav(id,btn){{
  document.querySelectorAll('.section').forEach(s=>s.classList.remove('active'));
  document.querySelectorAll('.sb-btn').forEach(b=>b.classList.remove('active'));
  document.getElementById(id).classList.add('active');
  btn.classList.add('active');
  document.getElementById('tb-title').textContent=TITLES[id]||id;
  document.querySelector('.content').scrollTop=0;
  closeSidebar();
}}
function openSidebar(){{document.getElementById('sidebar').classList.add('open');document.getElementById('sb-backdrop').classList.add('open');}}
function closeSidebar(){{document.getElementById('sidebar').classList.remove('open');document.getElementById('sb-backdrop').classList.remove('open');}}

document.addEventListener('DOMContentLoaded',function(){{
  document.getElementById('hb-btn').addEventListener('click',openSidebar);
  document.getElementById('sb-backdrop').addEventListener('click',closeSidebar);

  // Phase bars
  const phases={phases_js};
  const c=document.getElementById('phase-bars');
  phases.forEach(p=>{{
    const pct=((p.done/p.total)*100).toFixed(1);
    const r=document.createElement('div');r.className='ph-row';
    r.innerHTML=`<span class="ph-name">${{p.name}}</span><div class="ph-track"><div class="ph-bar" style="width:0%"></div></div><span class="ph-pct">${{pct}}%</span><span class="ph-cnt">${{p.done.toLocaleString()}} / ${{p.total.toLocaleString()}}</span>`;
    c.appendChild(r);
    setTimeout(()=>r.querySelector('.ph-bar').style.width=pct+'%',100);
  }});

  // Photos
  {photos_js}
  function mkCard(caption,src,idx){{
    const d=document.createElement('div');d.className='photo-card';
    d.innerHTML=`<img class="photo-img" src="${{src}}" alt="${{caption}}" loading="lazy"><div class="photo-cap">${{caption}}</div>`;
    d.querySelector('img').addEventListener('click',()=>openLB(idx));
    return d;
  }}
  {photo_render_js}

  // Lightbox
  let _idx=0;
  function openLB(idx){{_idx=idx;updateLB();document.getElementById('lb').classList.add('open');}}
  function updateLB(){{document.getElementById('lb-img').src=PHOTOS[_idx][1];document.getElementById('lb-cap').textContent=PHOTOS[_idx][0];}}
  document.getElementById('lb-close').onclick=()=>document.getElementById('lb').classList.remove('open');
  document.getElementById('lb-prev').onclick=()=>{{_idx=(_idx-1+PHOTOS.length)%PHOTOS.length;updateLB();}};
  document.getElementById('lb-next').onclick=()=>{{_idx=(_idx+1)%PHOTOS.length;updateLB();}};
  document.getElementById('lb').addEventListener('click',e=>{{if(e.target===e.currentTarget)e.currentTarget.classList.remove('open');}});
  document.addEventListener('keydown',e=>{{
    const lb=document.getElementById('lb');
    if(!lb.classList.contains('open'))return;
    if(e.key==='Escape')lb.classList.remove('open');
    if(e.key==='ArrowLeft'){{_idx=(_idx-1+PHOTOS.length)%PHOTOS.length;updateLB();}}
    if(e.key==='ArrowRight'){{_idx=(_idx+1)%PHOTOS.length;updateLB();}}
  }});
}});
</script>
{LOGIN_GATE_JS}
</body>
</html>"""

    return html


# ══════════════════════════════════════════════════════════════════════════════
# GIT PUSH
# ══════════════════════════════════════════════════════════════════════════════

def git_push(file_path, message=None):
    p = Path(file_path)
    if not message:
        m = re.search(r'IR-\d{3}', p.name)
        message = f"Weekly report {m.group(0) if m else 'IR-???'} — {datetime.now().strftime('%d.%m.%Y')}"
    print(f"\n🚀 Push su GitHub...")
    try:
        subprocess.run(['git','add',p.name],             check=True, capture_output=True, cwd=p.parent)
        subprocess.run(['git','commit','-m',message],    check=True, capture_output=True, cwd=p.parent)
        subprocess.run(['git','push','origin','main'],   check=True, capture_output=True, cwd=p.parent)
        print(f"  ✓ Pubblicato: {p.name}")
        return True
    except Exception as e:
        print(f"  ⚠️  {e}")
        print(f"  💡 Carica manualmente: {p.name}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════════════════════════════

HELP = """
EEW COMP3B — DOCX → HTML (full auto)
=====================================
Legge TUTTI i dati dal docx e genera la dashboard completa.

Uso:
  python auto_report.py <docx> [--output <dir>] [--push]

Esempi:
  python auto_report.py IR-076-EEW_COMP3B_-_18_05_2026_-_24_05_2026.docx
  python auto_report.py report.docx --output C:\\Reports --push

Dati letti automaticamente:
  ✓ Date visita, IR number, settimana
  ✓ Contatti (nome, azienda, ruolo, email)
  ✓ Dispatch per PO e categoria
  ✓ Pipes staccate ad Antwerp
  ✓ Fasi produzione (Phase Snapshot)
  ✓ Materiali e item summary
  ✓ Produzione annuale
  ✓ Attività settimanali (ITP table)
  ✓ VAGB plates KPI, delivery table, schedule
  ✓ HSE checklist
  ✓ Aree di concern (AOB)
  ✓ Foto con caption automatiche
  ✓ Enclosures (allegati)
"""

def main():
    if len(sys.argv) < 2 or sys.argv[1] in ('-h','--help'):
        print(HELP); return

    docx = sys.argv[1]
    push = '--push' in sys.argv
    out_dir = Path.cwd()
    if '--output' in sys.argv:
        i = sys.argv.index('--output')
        if i+1 < len(sys.argv): out_dir = Path(sys.argv[i+1])

    if not os.path.exists(docx):
        print(f"❌ File non trovato: {docx}"); sys.exit(1)

    print(f"\n{'='*58}")
    print(f"  EEW COMP3B  |  DOCX → HTML (full auto)")
    print(f"{'='*58}")

    # Parse
    parser = DocxParser(docx)
    d = parser.data

    # Build HTML
    print("\n🔨 Generazione HTML...")
    html = build_html(d)

    # Save
    out_dir.mkdir(parents=True, exist_ok=True)
    ir  = d['ir_number']
    out = out_dir / f"IR-{ir:03d}-EEW_COMP3B_redesign.html"
    out.write_text(html, encoding='utf-8')
    mb = out.stat().st_size / (1024*1024)

    print(f"\n✅ Completato!")
    print(f"   {out.name}  ({mb:.1f} MB)")

    if push:
        git_push(str(out))

if __name__ == '__main__':
    main()
